from docx import Document
import os
import subprocess  # For opening document on WSL Ubuntu
from docx.shared import Inches, Pt, RGBColor  # Styling headings
from docx.oxml.ns import qn  # Page numbers
from docx.oxml import OxmlElement  # Horizontal line, borders
from docx.enum.text import WD_ALIGN_PARAGRAPH  # For justification

from languages import languagesDataForDC, languageTranslations

leftLang = "english"
rightLang = "spanish"
numOfSections=138

document = Document()  # Create a new Word document

# Set smaller margins
sections = document.sections
for section in sections:
    section.top_margin = Inches(.5)  # Smaller margin for a more book-like appearance
    section.bottom_margin = Inches(.5)  # Smaller margin
    section.left_margin = Inches(.4)  # Smaller margin
    section.right_margin = Inches(.4)  # Smaller margin
    section.footer_distance = Inches(0.2)  # Make footer with page number smaller

def style_cell_text(cell, text, font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False):
    # Clear existing text
    cell.text = ''
    # Create a new run for the cell
    run = cell.paragraphs[0].add_run(text)
    # Apply the styles
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    # Set the alignment to justify
    cell.paragraphs[0].alignment = alignment
    # Adjust paragraph spacing
    paragraph_format = cell.paragraphs[0].paragraph_format
    paragraph_format.space_before = Pt(0)  # No space before paragraph
    paragraph_format.space_after = Pt(4)  # Small space after the paragraph
    paragraph_format.line_spacing = Pt(12)  # Adjusted line spacing

    # Apply cell borders
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'nil')
        b.set(qn('w:space'), '0')
        borders.append(b)

    # Set specific borders for the columns
    if cell._element.getparent().index(cell._element) % 2 == 0:  # First column
        right_border = OxmlElement('w:right')
        right_border.set(qn('w:val'), 'single')
        right_border.set(qn('w:sz'), '4')
        right_border.set(qn('w:space'), '0')
        borders.append(right_border)
    else:  # Second column
        left_border = OxmlElement('w:left')
        left_border.set(qn('w:val'), 'single')
        left_border.set(qn('w:sz'), '4')
        left_border.set(qn('w:space'), '0')
        borders.append(left_border)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    hr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Border size
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    hr.append(bottom)
    p._element.get_or_add_pPr().append(hr)

def add_title_page(doc):
    # Title page
    doc.add_paragraph("\n\n\n")  # Add spacing before the title
    # Add the main title in large, bold font
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
    run = main_title.add_run(f'{languagesDataForDC[leftLang]["dc"]}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(36)
    run.bold = True
    # Add spacing between title and subtitle
    doc.add_paragraph("\n\n")
    # Book of Mormon in second language
    second_title = doc.add_paragraph()
    second_title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
    run = second_title.add_run(f'{languagesDataForDC[rightLang]["dc"]}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(36)
    run.bold = True
    # Add more spacing before the side-by-side description
    doc.add_paragraph("\n\n")
    # Add a description below the titles
    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
    run = description.add_run(f'{languagesDataForDC[leftLang]["side-by-side-version"]}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True
    # Add language pairing description in smaller font
    language_pair = doc.add_paragraph()
    language_pair.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
    run = language_pair.add_run(f'{languageTranslations[leftLang][leftLang].capitalize()} | {languageTranslations[rightLang][rightLang].capitalize()}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    # Add more spacing before the page break
    doc.add_paragraph("\n\n\n")
    # Page break
    doc.add_page_break()

def create_introduction():
    add_horizontal_line(document)  # Line after book title
    document.add_paragraph("")  # Space after book title

    eng_path = f'dc2/dc-{leftLang}/{0}.txt' #will need to put dc/ before
    spa_path = f'dc2/dc-{rightLang}/{0}.txt'

    # Check if both files exist
    if os.path.exists(eng_path) and os.path.exists(spa_path):
        with open(eng_path, 'r', encoding='utf-8') as eng_file:
            english_verses = [line.strip() for line in eng_file.readlines() if line.strip()]  # Removes new line characters

        with open(spa_path, 'r', encoding='utf-8') as spa_file:
            spanish_verses = [line.strip() for line in spa_file.readlines() if line.strip()]  # Removes new line characters

        # Create a table with two columns
        table = document.add_table(rows=0, cols=2)

        # Add borders to the columns to create a vertical line between them
        for i, cell in enumerate(table.columns[0].cells + table.columns[1].cells):
            if cell._element.getparent().index(cell._element) % 2 == 0:  # First column
                cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                borders = cell._element.find(qn('w:tcBorders'))
                if borders is None:
                    borders = OxmlElement('w:tcBorders')
                    cell._element.get_or_add_tcPr().append(borders)
                right_border = OxmlElement('w:right')
                right_border.set(qn('w:val'), 'single')
                right_border.set(qn('w:sz'), '4')
                right_border.set(qn('w:space'), '0')
                borders.append(right_border)
            else:  # Second column
                cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                borders = cell._element.find(qn('w:tcBorders'))
                if borders is None:
                    borders = OxmlElement('w:tcBorders')
                    cell._element.get_or_add_tcPr().append(borders)
                left_border = OxmlElement('w:left')
                left_border.set(qn('w:val'), 'single')
                left_border.set(qn('w:sz'), '4')
                left_border.set(qn('w:space'), '0')
                borders.append(left_border)

        

            add_horizontal_line(document)  # Line after chapter
        
        #introduction
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[0].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        style_cell_text(row_cells[1], f"{spanish_verses[0].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        #space
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
        style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
        
        # Add paragraphs without verse numvers
        for i in range(1, 12): #NOTE that english verses should be the same length as spanish verses
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"    {english_verses[i].strip()}")
            style_cell_text(row_cells[1], f"    {spanish_verses[i].strip()}")
        #space
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
        style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
        #Testimony of the Twelve Apostles to the Truth of the Book of Doctrine and Covenants
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[13].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        style_cell_text(row_cells[1], f"{spanish_verses[13].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        #italiced paragraph
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"    {english_verses[14]}", font_name='Times New Roman', font_size=12, italic=True)
        style_cell_text(row_cells[1], f"    {spanish_verses[14]}", font_name='Times New Roman', font_size=12, italic=True)
        # Add paragraphs without verse numvers
        for i in range(14, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"    {english_verses[i].strip()}")
            style_cell_text(row_cells[1], f"    {spanish_verses[i].strip()}")

def create_sections():
    # Iterate through each section
    for chapter in range(1,numOfSections+1):
        add_horizontal_line(document)  # Line after book title

        eng_path = f'dc2/dc-{leftLang}/{chapter}.txt'
        spa_path = f'dc2/dc-{rightLang}/{chapter}.txt'
        
        # Check if both files exist
        if os.path.exists(eng_path) and os.path.exists(spa_path):
            with open(eng_path, 'r', encoding='utf-8') as eng_file:
                english_verses = [line.strip() for line in eng_file.readlines() if line.strip()]  # Removes new line characters

            with open(spa_path, 'r', encoding='utf-8') as spa_file:
                spanish_verses = [line.strip() for line in spa_file.readlines() if line.strip()]  # Removes new line characters

            # Create a table with two columns
            table = document.add_table(rows=0, cols=2)

            # Add borders to the columns to create a vertical line between them
            for i, cell in enumerate(table.columns[0].cells + table.columns[1].cells):
                if cell._element.getparent().index(cell._element) % 2 == 0:  # First column
                    cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                    borders = cell._element.find(qn('w:tcBorders'))
                    if borders is None:
                        borders = OxmlElement('w:tcBorders')
                        cell._element.get_or_add_tcPr().append(borders)
                    right_border = OxmlElement('w:right')
                    right_border.set(qn('w:val'), 'single')
                    right_border.set(qn('w:sz'), '4')
                    right_border.set(qn('w:space'), '0')
                    borders.append(right_border)
                else:  # Second column
                    cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                    borders = cell._element.find(qn('w:tcBorders'))
                    if borders is None:
                        borders = OxmlElement('w:tcBorders')
                        cell._element.get_or_add_tcPr().append(borders)
                    left_border = OxmlElement('w:left')
                    left_border.set(qn('w:val'), 'single')
                    left_border.set(qn('w:sz'), '4')
                    left_border.set(qn('w:space'), '0')
                    borders.append(left_border)

        

                add_horizontal_line(document)  # Line after chapter
            
            #CHAPTER 1
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"{languagesDataForDC[leftLang]['section'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
            style_cell_text(row_cells[1], f"{languagesDataForDC[rightLang]['section'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
            # heading
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
            style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
            # breakdown
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"{english_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)
            style_cell_text(row_cells[1], f"{spanish_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)
            #space
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
            style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
            
            # Add verses to the table with verse numbers
            for i in range(2, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], f"{i-1} {english_verses[i].strip()}")
                style_cell_text(row_cells[1], f"{i-1} {spanish_verses[i].strip()}")

def create_od1():
    add_horizontal_line(document)  # Line after book title
    document.add_paragraph("")  # Space after book title

    eng_path = f'dc2/dc-{leftLang}/od{1}.txt' #will need to put dc/ before
    spa_path = f'dc2/dc-{rightLang}/od{1}.txt'

    # Check if both files exist
    if os.path.exists(eng_path) and os.path.exists(spa_path):
        with open(eng_path, 'r', encoding='utf-8') as eng_file:
            english_verses = [line.strip() for line in eng_file.readlines() if line.strip()]  # Removes new line characters

        with open(spa_path, 'r', encoding='utf-8') as spa_file:
            spanish_verses = [line.strip() for line in spa_file.readlines() if line.strip()]  # Removes new line characters

        # Create a table with two columns
        table = document.add_table(rows=0, cols=2)

        # Add borders to the columns to create a vertical line between them
        for i, cell in enumerate(table.columns[0].cells + table.columns[1].cells):
            if cell._element.getparent().index(cell._element) % 2 == 0:  # First column
                cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                borders = cell._element.find(qn('w:tcBorders'))
                if borders is None:
                    borders = OxmlElement('w:tcBorders')
                    cell._element.get_or_add_tcPr().append(borders)
                right_border = OxmlElement('w:right')
                right_border.set(qn('w:val'), 'single')
                right_border.set(qn('w:sz'), '4')
                right_border.set(qn('w:space'), '0')
                borders.append(right_border)
            else:  # Second column
                cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                borders = cell._element.find(qn('w:tcBorders'))
                if borders is None:
                    borders = OxmlElement('w:tcBorders')
                    cell._element.get_or_add_tcPr().append(borders)
                left_border = OxmlElement('w:left')
                left_border.set(qn('w:val'), 'single')
                left_border.set(qn('w:sz'), '4')
                left_border.set(qn('w:space'), '0')
                borders.append(left_border)

        

            add_horizontal_line(document)  # Line after chapter
        
        #offical declatation 1
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[0].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        style_cell_text(row_cells[1], f"{spanish_verses[0].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        #space
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
        style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
        #italics
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[1]}", font_name='Times New Roman', font_size=12, italic=True)
        style_cell_text(row_cells[1], f"{spanish_verses[1]}", font_name='Times New Roman', font_size=12, italic=True)
        #space
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
        style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
        #to whom it may concern
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[2]}", font_name='Times New Roman', font_size=12)
        style_cell_text(row_cells[1], f"{spanish_verses[2]}", font_name='Times New Roman', font_size=12)
        # Add paragraphs without verse numvers
        for i in range(3,8): #NOTE that english verses should be the same length as spanish verses
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"    {english_verses[i].strip()}")
            style_cell_text(row_cells[1], f"    {spanish_verses[i].strip()}")
        #Wilford Woodrfu
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[8].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        style_cell_text(row_cells[1], f"{spanish_verses[8].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        #Prez of church
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[9]}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        style_cell_text(row_cells[1], f"{spanish_verses[9]}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        #space
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
        style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
        # Add paragraphs without verse numvers
        for i in range(10,12): #NOTE that english verses should be the same length as spanish verses
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"    {english_verses[i].strip()}")
            style_cell_text(row_cells[1], f"    {spanish_verses[i].strip()}")
        #date
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[12]}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        style_cell_text(row_cells[1], f"{spanish_verses[12]}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        #space
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
        style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
        #exceprts
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[13].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        style_cell_text(row_cells[1], f"{spanish_verses[13].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        # Add paragraphs without verse numvers
        for i in range(14,len(english_verses)): #NOTE that english verses should be the same length as spanish verses
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"    {english_verses[i].strip()}")
            style_cell_text(row_cells[1], f"    {spanish_verses[i].strip()}")
        
def create_od2():
    add_horizontal_line(document)  # Line after book title
    document.add_paragraph("")  # Space after book title

    eng_path = f'dc2/dc-{leftLang}/od{2}.txt' #will need to put dc/ before
    spa_path = f'dc2/dc-{rightLang}/od{2}.txt'

    # Check if both files exist
    if os.path.exists(eng_path) and os.path.exists(spa_path):
        with open(eng_path, 'r', encoding='utf-8') as eng_file:
            english_verses = [line.strip() for line in eng_file.readlines() if line.strip()]  # Removes new line characters

        with open(spa_path, 'r', encoding='utf-8') as spa_file:
            spanish_verses = [line.strip() for line in spa_file.readlines() if line.strip()]  # Removes new line characters

        # Create a table with two columns
        table = document.add_table(rows=0, cols=2)

        # Add borders to the columns to create a vertical line between them
        for i, cell in enumerate(table.columns[0].cells + table.columns[1].cells):
            if cell._element.getparent().index(cell._element) % 2 == 0:  # First column
                cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                borders = cell._element.find(qn('w:tcBorders'))
                if borders is None:
                    borders = OxmlElement('w:tcBorders')
                    cell._element.get_or_add_tcPr().append(borders)
                right_border = OxmlElement('w:right')
                right_border.set(qn('w:val'), 'single')
                right_border.set(qn('w:sz'), '4')
                right_border.set(qn('w:space'), '0')
                borders.append(right_border)
            else:  # Second column
                cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                borders = cell._element.find(qn('w:tcBorders'))
                if borders is None:
                    borders = OxmlElement('w:tcBorders')
                    cell._element.get_or_add_tcPr().append(borders)
                left_border = OxmlElement('w:left')
                left_border.set(qn('w:val'), 'single')
                left_border.set(qn('w:sz'), '4')
                left_border.set(qn('w:space'), '0')
                borders.append(left_border)

        

            add_horizontal_line(document)  # Line after chapter
        
        #offical declatation 2
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[0].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        style_cell_text(row_cells[1], f"{spanish_verses[0].upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        #space
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
        style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
        #italics
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[1]}", font_name='Times New Roman', font_size=12, italic=True)
        style_cell_text(row_cells[1], f"{spanish_verses[1]}", font_name='Times New Roman', font_size=12, italic=True)
        #space
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
        style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
        #to whom it may concern
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[2]}", font_name='Times New Roman', font_size=12)
        style_cell_text(row_cells[1], f"{spanish_verses[2]}", font_name='Times New Roman', font_size=12)
        # Add paragraphs with indentation
        for i in range(3,6): #NOTE that english verses should be the same length as spanish verses
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"    {english_verses[i].strip()}")
            style_cell_text(row_cells[1], f"    {spanish_verses[i].strip()}")
        #space
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
        style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
        # Add paragraphs without indentation
        for i in range(6,9): #NOTE that english verses should be the same length as spanish verses
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"{english_verses[i].strip()}")
            style_cell_text(row_cells[1], f"{spanish_verses[i].strip()}")
        # Add paragraphs with indentation
        for i in range(9,13): #NOTE that english verses should be the same length as spanish verses
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"    {english_verses[i].strip()}")
            style_cell_text(row_cells[1], f"    {spanish_verses[i].strip()}")
        #sincerly yours
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"    {english_verses[13]}", font_name='Times New Roman', font_size=12)
        style_cell_text(row_cells[1], f"    {spanish_verses[13]}", font_name='Times New Roman', font_size=12)
        # first prez names
        for i in range(14,17): #NOTE that english verses should be the same length as spanish verses
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"    {english_verses[i].upper()}")
            style_cell_text(row_cells[1], f"    {spanish_verses[i].upper()}")
        #the first presidency
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"    {english_verses[17]}")
        style_cell_text(row_cells[1], f"    {spanish_verses[17]}")
        #space
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
        style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
        # Add paragraphs without verse numvers
        for i in range(18,20): #NOTE that english verses should be the same length as spanish verses
            row_cells = table.add_row().cells
            style_cell_text(row_cells[0], f"    {english_verses[i].strip()}")
            style_cell_text(row_cells[1], f"    {spanish_verses[i].strip()}")
        #date
        row_cells = table.add_row().cells
        style_cell_text(row_cells[0], f"{english_verses[20]}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        style_cell_text(row_cells[1], f"{spanish_verses[20]}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
               

# Page numbers
def add_page_numbers(document):
    sections = document.sections
    for section in sections:
        footer = section.footer
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the page number
        run = p.add_run()
        field = OxmlElement('w:fldSimple')
        field.set(qn('w:instr'), 'PAGE')  # PAGE is the instruction for page number
        run._element.append(field)


add_title_page(document) 
create_introduction()
create_sections()
create_od1()
create_od2()
add_page_numbers(document)

# Save the document
document.save("side_by_side_dc.docx")
subprocess.call(['powershell.exe', 'Start-Process', 'side_by_side_dc.docx'])  # Opens document
