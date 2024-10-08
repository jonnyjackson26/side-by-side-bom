from docx import Document
import os
import subprocess  # For opening document on WSL Ubuntu
from docx.shared import Inches, Pt, RGBColor  # Styling headings
from docx.oxml.ns import qn  # Page numbers
from docx.oxml import OxmlElement  # Horizontal line, borders
from docx.enum.text import WD_ALIGN_PARAGRAPH  # For justification

from languages import languagesData

leftLang = "portuguese"
rightLang = "spanish"

books = ["title","introduction", "three","eight","js", "1-nephi", "2-nephi", "jacob", "enos", "jarom", "omni", "words-of-mormon", "mosiah", "alma", "helaman", "3-nephi", "4-nephi", "mormon", "ether", "moroni"]
chapters = {
    "title":1,"introduction":1,"three":1,"eight":1,"js":1,
    "1-nephi": 22, "2-nephi": 33, "jacob": 7, "enos": 1, "jarom": 1, "omni": 1, "words-of-mormon": 1,
    "mosiah": 29, "alma": 63, "helaman": 16, "3-nephi": 30, "4-nephi": 1,
    "mormon": 9, "ether": 15, "moroni": 10
}

document = Document()  # Create a new Word document

# Set smaller margins
sections = document.sections
for section in sections:
    section.top_margin = Inches(.5)  # Smaller margin for a more book-like appearance
    section.bottom_margin = Inches(.5)  # Smaller margin
    section.left_margin = Inches(.4)  # Smaller margin
    section.right_margin = Inches(.4)  # Smaller margin
    section.footer_distance = Inches(0.2)  # Make footer with page number smaller

def style_cell_text(cell, text, font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False):
    # Clear existing text
    cell.text = ''
    # Create a new run for the cell
    run = cell.paragraphs[0].add_run(text)
    # Apply the styles
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    # Set the alignment to justify
    cell.paragraphs[0].alignment = alignment
    # Adjust paragraph spacing
    paragraph_format = cell.paragraphs[0].paragraph_format
    paragraph_format.space_before = Pt(0)  # No space before paragraph
    paragraph_format.space_after = Pt(4)  # Small space after the paragraph
    paragraph_format.line_spacing = Pt(12)  # Adjusted line spacing

    # Apply cell borders
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'nil')
        b.set(qn('w:space'), '0')
        borders.append(b)

    # Set specific borders for the columns
    if cell._element.getparent().index(cell._element) % 2 == 0:  # First column
        right_border = OxmlElement('w:right')
        right_border.set(qn('w:val'), 'single')
        right_border.set(qn('w:sz'), '4')
        right_border.set(qn('w:space'), '0')
        borders.append(right_border)
    else:  # Second column
        left_border = OxmlElement('w:left')
        left_border.set(qn('w:val'), 'single')
        left_border.set(qn('w:sz'), '4')
        left_border.set(qn('w:space'), '0')
        borders.append(left_border)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    hr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Border size
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    hr.append(bottom)
    p._element.get_or_add_pPr().append(hr)

def add_title_page(doc):
    # Title page
    doc.add_paragraph("\n\n\n")  # Add spacing before the title
    # Add the main title in large, bold font
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
    run = main_title.add_run(f'{languagesData[leftLang]["book-of-mormon"]}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(36)
    run.bold = True
    # Add subtitle in a slightly smaller font and italics
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
    run = subtitle.add_run(f'{languagesData[leftLang]["another-testament-of-jesus-christ"]}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.italic = True
    # Add spacing between title and subtitle
    doc.add_paragraph("\n\n")
    # Book of Mormon in second language
    second_title = doc.add_paragraph()
    second_title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
    run = second_title.add_run(f'{languagesData[rightLang]["book-of-mormon"]}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(36)
    run.bold = True
    # Add another subtitle for the translation language in a smaller font
    subtitle_3 = doc.add_paragraph()
    subtitle_3.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
    run = subtitle_3.add_run(f'{languagesData[rightLang]["another-testament-of-jesus-christ"]}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.italic = True
    # Add more spacing before the side-by-side description
    doc.add_paragraph("\n\n")
    # Add a description below the titles
    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
    run = description.add_run(f'{languagesData[leftLang]["side-by-side-version"]}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True
    # Add language pairing description in smaller font
    language_pair = doc.add_paragraph()
    language_pair.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
    run = language_pair.add_run(f'{languagesData[leftLang][leftLang].capitalize()} | {languagesData[rightLang][rightLang].capitalize()}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    # Add more spacing before the page break
    doc.add_paragraph("\n\n\n")
    # Page break
    doc.add_page_break()

add_title_page(document) # Add title page

# Iterate through each book
for book in books:
    add_horizontal_line(document)  # Line after book title
    document.add_paragraph("")  # Space after book title

    # Iterate through each chapter
    for chapter in range(1, chapters[book] + 1):
        eng_path = f'bom3/bom-{leftLang}/{book}/{chapter}.txt'
        spa_path = f'bom3/bom-{rightLang}/{book}/{chapter}.txt'
        
        # Check if both files exist
        if os.path.exists(eng_path) and os.path.exists(spa_path):
            with open(eng_path, 'r', encoding='utf-8') as eng_file:
                english_verses = [line.strip() for line in eng_file.readlines() if line.strip()]  # Removes new line characters

            with open(spa_path, 'r', encoding='utf-8') as spa_file:
                spanish_verses = [line.strip() for line in spa_file.readlines() if line.strip()]  # Removes new line characters

            # Create a table with two columns
            table = document.add_table(rows=0, cols=2)

            # Add borders to the columns to create a vertical line between them
            for i, cell in enumerate(table.columns[0].cells + table.columns[1].cells):
                if cell._element.getparent().index(cell._element) % 2 == 0:  # First column
                    cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                    borders = cell._element.find(qn('w:tcBorders'))
                    if borders is None:
                        borders = OxmlElement('w:tcBorders')
                        cell._element.get_or_add_tcPr().append(borders)
                    right_border = OxmlElement('w:right')
                    right_border.set(qn('w:val'), 'single')
                    right_border.set(qn('w:sz'), '4')
                    right_border.set(qn('w:space'), '0')
                    borders.append(right_border)
                else:  # Second column
                    cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))
                    borders = cell._element.find(qn('w:tcBorders'))
                    if borders is None:
                        borders = OxmlElement('w:tcBorders')
                        cell._element.get_or_add_tcPr().append(borders)
                    left_border = OxmlElement('w:left')
                    left_border.set(qn('w:val'), 'single')
                    left_border.set(qn('w:sz'), '4')
                    left_border.set(qn('w:space'), '0')
                    borders.append(left_border)

            
            if book=="1-nephi": #with book intros and also like "his reign and ministry" or something after the book title
                if chapter==1:
                    # The First Book of Nephi
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    # His Reign and Ministry
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[1].strip().upper()}", font_name='Times New Roman', font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_text(row_cells[1], f"{spanish_verses[1].strip().upper()}", font_name='Times New Roman', font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    # BOOK INTRO (An account of Lehi and his wife Sariah, a...)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[2].strip()}", font_name='Times New Roman', font_size=12)
                    style_cell_text(row_cells[1], f"{spanish_verses[2].strip()}", font_name='Times New Roman', font_size=12)
                    #space
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                    style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    #CHAPTER 1
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    #chapter heading 
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[3].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[3].strip()}", font_name='Times New Roman', font_size=12, italic=True)


                    # Add verses to the table with verse numbers
                    for i in range(4, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i-3} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i-3} {spanish_verses[i].strip()}")

                else: #normal chapter with just one line of chapter intro and the other lines are verses
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(1, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i} {spanish_verses[i].strip()}")
            if book=="4-nephi": #with book intros and also like "his reign and ministry" or something after the book title and then the book intro is in italics and centerd
                if chapter==1:
                    # The First Book of Nephi
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    # His Reign and Ministry
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[1].strip().upper()}", font_name='Times New Roman', font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_text(row_cells[1], f"{spanish_verses[1].strip().upper()}", font_name='Times New Roman', font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    # BOOK INTRO (An account of Lehi and his wife Sariah, a...)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[2].strip()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[2].strip()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
                    #space
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                    style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    #CHAPTER 1
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    #chapter heading 
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[3].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[3].strip()}", font_name='Times New Roman', font_size=12, italic=True)


                    # Add verses to the table with verse numbers
                    for i in range(4, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i-3} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i-3} {spanish_verses[i].strip()}")

                else: #normal chapter with just one line of chapter intro and the other lines are verses
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(1, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i} {spanish_verses[i].strip()}")
            elif book=="2-nephi" or book=="jacob" or book=="ether": #with book intros
                if chapter==1:
                    # The Second Book of Nephi
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    # BOOK INTRO (An account of Lehi and his wife Sariah, a...)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[1].strip()}", font_name='Times New Roman', font_size=12)
                    style_cell_text(row_cells[1], f"{spanish_verses[1].strip()}", font_name='Times New Roman', font_size=12)
                    #space
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                    style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    #CHAPTER 1
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    #chapter heading 
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)


                    # Add verses to the table with verse numbers
                    for i in range(3, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i-2} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i-2} {spanish_verses[i].strip()}")

                else: #normal chapter with just one line of chapter intro and the other lines are verses
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(1, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i} {spanish_verses[i].strip()}")
            elif book=="helaman": #with book intro but no weird little "his reign and ministry"
                if chapter==1:
                    # The Second Book of Nephi
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    #space
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                    style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    #CHAPTER 1
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    #chapter heading 
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)


                    # Add verses to the table with verse numbers
                    for i in range(3, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i-2} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i-2} {spanish_verses[i].strip()}")
                elif chapter==7 or chapter==13:
                    # the words of alma...
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    # comprising ch7
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    #space
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                    style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(3, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i-2} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i-2} {spanish_verses[i].strip()}")

                else: #normal chapter with just one line of chapter intro and the other lines are verses
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(1, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i} {spanish_verses[i].strip()}")
            if book=="alma": #no book intros but "his reign and ministry" or something after the book title
                if chapter==1:
                    # The First Book of Nephi
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    # book intro
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[1].strip()}", font_name='Times New Roman', font_size=12)
                    style_cell_text(row_cells[1], f"{spanish_verses[1].strip()}", font_name='Times New Roman', font_size=12)
                    #space
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                    style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    #CHAPTER 1
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    #chapter heading 
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)


                    # Add verses to the table with verse numbers
                    for i in range(3, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i-2} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i-2} {spanish_verses[i].strip()}")
                elif chapter==7 or chapter==5 or chapter==36 or chapter==9 or chapter==45 or chapter==39 or chapter==38 or chapter==21 or chapter==17: #two line of italics before
                    # the words of alma...
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    # comprising ch7
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    #space
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                    style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(3, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i-2} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i-2} {spanish_verses[i].strip()}")
                else: #normal chapter with just one line of chapter intro and the other lines are verses
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(1, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i} {spanish_verses[i].strip()}")
            if book=="3-nephi": #without book intros but like "his reign and ministry" and then a smaller like "who was the son of ..." book title
                if chapter==1:
                    # The First Book of Nephi
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    # His Reign and Ministry
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[1].strip().upper()}", font_name='Times New Roman', font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_text(row_cells[1], f"{spanish_verses[1].strip().upper()}", font_name='Times New Roman', font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    # BOOK INTRO (An account of Lehi and his wife Sariah, a...)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    #space
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                    style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    #CHAPTER 1
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    #chapter heading 
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[3].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[3].strip()}", font_name='Times New Roman', font_size=12, italic=True)


                    # Add verses to the table with verse numbers
                    for i in range(4, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i-3} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i-3} {spanish_verses[i].strip()}")
                elif chapter==11: #two lines of italics
                    # jesus christ did show ...
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    # comprising ch7
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    #space
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                    style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(3, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i-2} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i-2} {spanish_verses[i].strip()}")

                else: #normal chapter with just one line of chapter intro and the other lines are verses
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(1, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i} {spanish_verses[i].strip()}")
            if book=="enos" or book=="jarom" or book=="omni" or book=="words-of-mormon" or book=="mosiah" or book=="mormon" or book=="moroni": #no intro, just bookname and then chapter heading
                if chapter==1:
                    # The First Book of Nephi
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    #space
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                    style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    #CHAPTER 1
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    #chapter heading 
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)


                    # Add verses to the table with verse numbers
                    for i in range(3, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i-2} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i-2} {spanish_verses[i].strip()}")
                elif book=="moroni" and chapter==10: #the end text
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(1, len(english_verses)-1): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i} {spanish_verses[i].strip()}")

                    #the end text
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[-1].strip().upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_text(row_cells[1], f"{spanish_verses[-1].strip().upper()}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                elif (book=="mosiah" and (chapter==9 or chapter==23)) or (book=="moroni" and chapter==9): #two line of italics before
                    # the words of alma...
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    # comprising ch7
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[1].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    #space
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                    style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[2].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(3, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i-2} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i-2} {spanish_verses[i].strip()}")
                
                else: #normal chapter with just one line of chapter intro and the other lines are verses
                    # Have the first row of the columns be "Chapter X" and "Capítulo X"
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{languagesData[leftLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    style_cell_text(row_cells[1], f"{languagesData[rightLang]['chapter'].upper()} {chapter}", font_name='Times New Roman', font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

                    # Add chapter headings (the 0th index of the verses list is added with italics)
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)
                    style_cell_text(row_cells[1], f"{spanish_verses[0].strip()}", font_name='Times New Roman', font_size=12, italic=True)

                    # Add verses to the table with verse numbers
                    for i in range(1, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                        row_cells = table.add_row().cells
                        style_cell_text(row_cells[0], f"{i} {english_verses[i].strip()}")
                        style_cell_text(row_cells[1], f"{i} {spanish_verses[i].strip()}")
            if book=="title":
                 # the book fo mormon
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], f"{english_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                style_cell_text(row_cells[1], f"{spanish_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                # account taken from hand of mormon
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], f"{english_verses[1].strip().upper()}", font_name='Times New Roman', font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                style_cell_text(row_cells[1], f"{spanish_verses[1].strip().upper()}", font_name='Times New Roman', font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                #space
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                    
                for i in range(2, len(english_verses)-1): #NOTE that english verses should be the same length as spanish verses
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"  {english_verses[i].strip()}")
                    style_cell_text(row_cells[1], f"  {spanish_verses[i].strip()}")
                # translated by josepg smith jr
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], f"{english_verses[-1].strip().upper()}", font_name='Times New Roman', font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                style_cell_text(row_cells[1], f"{spanish_verses[-1].strip().upper()}", font_name='Times New Roman', font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            if book=="introduction":
                 # introduction
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], f"{english_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                style_cell_text(row_cells[1], f"{spanish_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                #space
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                for i in range(1, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"    {english_verses[i]}")
                    style_cell_text(row_cells[1], f"    {spanish_verses[i]}")
            if book=="three":
                 # three witnesses
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], f"{english_verses[0].upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                style_cell_text(row_cells[1], f"{spanish_verses[0].upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                #space
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                #their testimony
                for i in range(1, len(english_verses)-3): #NOTE that english verses should be the same length as spanish verses
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"    {english_verses[i]}")
                    style_cell_text(row_cells[1], f"    {spanish_verses[i]}")
                #each of their names
                for i in range(len(english_verses)-3, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[i]}", alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                    style_cell_text(row_cells[1], f"{spanish_verses[i]}", alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            if book=="eight":
                 # eight witnesses
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], f"{english_verses[0].upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                style_cell_text(row_cells[1], f"{spanish_verses[0].upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                #space
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                #their testimony
                for i in range(1, len(english_verses)-8): #NOTE that english verses should be the same length as spanish verses
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"    {english_verses[i]}")
                    style_cell_text(row_cells[1], f"    {spanish_verses[i]}")
                #each of their names
                for i in range(len(english_verses)-8, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"{english_verses[i]}", alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                    style_cell_text(row_cells[1], f"{spanish_verses[i]}", alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            if book=="js":
                # testimony of prophet josehp
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], f"{english_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                style_cell_text(row_cells[1], f"{spanish_verses[0].strip().upper()}", font_name='Times New Roman', font_size=15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                #space
                row_cells = table.add_row().cells
                style_cell_text(row_cells[0], "", font_name='Times New Roman', font_size=5)
                style_cell_text(row_cells[1], "", font_name='Times New Roman', font_size=5)
                for i in range(1, len(english_verses)): #NOTE that english verses should be the same length as spanish verses
                    row_cells = table.add_row().cells
                    style_cell_text(row_cells[0], f"    {english_verses[i]}")
                    style_cell_text(row_cells[1], f"    {spanish_verses[i]}")


            add_horizontal_line(document)  # Line after chapter
        

# Page numbers
def add_page_numbers(document):
    sections = document.sections
    for section in sections:
        footer = section.footer
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the page number
        run = p.add_run()
        field = OxmlElement('w:fldSimple')
        field.set(qn('w:instr'), 'PAGE')  # PAGE is the instruction for page number
        run._element.append(field)

# Add page numbers to the footer
add_page_numbers(document)

# Save the document
document.save("side_by_side_bom.docx")
subprocess.call(['powershell.exe', 'Start-Process', 'side_by_side_bom.docx'])  # Opens document
